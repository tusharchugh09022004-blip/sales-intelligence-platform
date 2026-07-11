from datetime import datetime
from pathlib import Path

import pandas as pd
import matplotlib
matplotlib.use("Agg")
import matplotlib.pyplot as plt
from docx import Document
from docx.enum.section import WD_SECTION_START
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parent
DATASET = ROOT / "dataset" / "enriched_sales_dataset.csv"
OUT = ROOT / "PTU_Industrial_Training_Report_Sales_Intelligence_Platform.docx"
FIG_DIR = ROOT / "report_figures"


STUDENT_NAME = "[NAME OF STUDENT]"
ROLL_NO = "[UNIVERSITY ROLL NO.]"
COMPANY = "[NAME OF COMPANY / INDUSTRY / INSTITUTE]"
PROGRAM = "B.Tech. Computer Science and Engineering"
TRAINING_PERIOD = "January 2025 - May 2025"


def set_run_font(run, size=12, bold=False, italic=False, color=None):
    run.font.name = "Times New Roman"
    run._element.rPr.rFonts.set(qn("w:ascii"), "Times New Roman")
    run._element.rPr.rFonts.set(qn("w:hAnsi"), "Times New Roman")
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic
    if color:
        run.font.color.rgb = RGBColor.from_string(color)


def set_cell_text(cell, text, size=10, bold=False, align=WD_ALIGN_PARAGRAPH.LEFT):
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    p = cell.paragraphs[0]
    p.alignment = align
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.line_spacing = 1
    p.text = ""
    r = p.add_run(str(text))
    set_run_font(r, size=size, bold=bold)


def add_page_field(paragraph):
    run = paragraph.add_run()
    fld_char1 = OxmlElement("w:fldChar")
    fld_char1.set(qn("w:fldCharType"), "begin")
    instr_text = OxmlElement("w:instrText")
    instr_text.set(qn("xml:space"), "preserve")
    instr_text.text = "PAGE"
    fld_char2 = OxmlElement("w:fldChar")
    fld_char2.set(qn("w:fldCharType"), "end")
    run._r.append(fld_char1)
    run._r.append(instr_text)
    run._r.append(fld_char2)


def configure_document(doc):
    section = doc.sections[0]
    section.page_width = Cm(21.0)
    section.page_height = Cm(29.7)
    section.left_margin = Cm(3.5)
    section.right_margin = Cm(1.25)
    section.top_margin = Cm(2.5)
    section.bottom_margin = Cm(1.25)
    section.header_distance = Cm(1.25)
    section.footer_distance = Cm(0.8)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Times New Roman"
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Times New Roman")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Times New Roman")
    normal.font.size = Pt(12)
    normal.paragraph_format.line_spacing = 2
    normal.paragraph_format.space_after = Pt(0)

    for name, size in [("Heading 1", 16), ("Heading 2", 14), ("Heading 3", 12)]:
        style = styles[name]
        style.font.name = "Times New Roman"
        style._element.rPr.rFonts.set(qn("w:ascii"), "Times New Roman")
        style._element.rPr.rFonts.set(qn("w:hAnsi"), "Times New Roman")
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor(0, 0, 0)
        style.paragraph_format.line_spacing = 2
        style.paragraph_format.space_before = Pt(6)
        style.paragraph_format.space_after = Pt(6)

    header = section.header.paragraphs[0]
    header.text = ""
    header.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = header.add_run("Dept. of CSE, IKG Punjab Technical University, Kapurthala")
    set_run_font(r, size=10, italic=True)

    footer = section.footer.paragraphs[0]
    footer.text = ""
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = footer.add_run(f"{ROLL_NO}        ")
    set_run_font(r, size=10)
    add_page_field(footer)


def para(doc, text="", size=12, bold=False, italic=False, align=None, spacing=2):
    p = doc.add_paragraph()
    if align is not None:
        p.alignment = align
    p.paragraph_format.line_spacing = spacing
    p.paragraph_format.space_after = Pt(0)
    r = p.add_run(text)
    set_run_font(r, size=size, bold=bold, italic=italic)
    return p


def chapter(doc, title):
    doc.add_page_break()
    p = para(doc, title.upper(), size=16, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)
    p.paragraph_format.space_after = Pt(12)


def heading(doc, text, level=2):
    p = doc.add_paragraph(text, style=f"Heading {level}")
    if level == 1:
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    return p


def bullet(doc, text):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.line_spacing = 2
    p.paragraph_format.space_after = Pt(0)
    r = p.add_run(text)
    set_run_font(r)


def number_item(doc, text):
    p = doc.add_paragraph(style="List Number")
    p.paragraph_format.line_spacing = 2
    p.paragraph_format.space_after = Pt(0)
    r = p.add_run(text)
    set_run_font(r)


def caption(doc, text, kind="figure"):
    p = para(doc, text, size=10, italic=True, align=WD_ALIGN_PARAGRAPH.CENTER, spacing=1)
    p.paragraph_format.space_before = Pt(3)
    p.paragraph_format.space_after = Pt(6)
    return p


def table_caption(doc, text):
    p = para(doc, text, size=10, italic=True, align=WD_ALIGN_PARAGRAPH.CENTER, spacing=1)
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(3)
    return p


def add_table(doc, rows, widths=None):
    table = doc.add_table(rows=1, cols=len(rows[0]))
    table.style = "Table Grid"
    hdr = table.rows[0].cells
    for i, value in enumerate(rows[0]):
        set_cell_text(hdr[i], value, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)
    for row in rows[1:]:
        cells = table.add_row().cells
        for i, value in enumerate(row):
            align = WD_ALIGN_PARAGRAPH.CENTER if len(str(value)) < 18 else WD_ALIGN_PARAGRAPH.LEFT
            set_cell_text(cells[i], value, align=align)
    if widths:
        for row in table.rows:
            for idx, width in enumerate(widths):
                row.cells[idx].width = Cm(width)
    doc.add_paragraph()
    return table


def load_stats():
    df = pd.read_csv(DATASET)
    df["Order Date"] = pd.to_datetime(df["Order Date"], errors="coerce")
    stats = {
        "rows": len(df),
        "orders": df["Order ID"].nunique(),
        "customers": df["Customer Name"].nunique(),
        "products": df["Product Name"].nunique(),
        "sales": df["Sales Amount"].sum(),
        "profit": df["Profit"].sum(),
        "margin": df["Profit"].sum() / df["Sales Amount"].sum() * 100,
        "avg_discount": df["Discount"].mean() * 100,
        "years": ", ".join(map(str, sorted(df["Order Date"].dt.year.dropna().unique().astype(int)))),
        "regions": ", ".join(sorted(df["Region"].dropna().unique())),
        "categories": ", ".join(sorted(df["Product Category"].dropna().unique())),
        "columns": len(df.columns),
    }
    return df, stats


def make_figures(df):
    FIG_DIR.mkdir(exist_ok=True)
    figs = {}

    region = df.groupby("Region")["Sales Amount"].sum().sort_values(ascending=False)
    plt.figure(figsize=(7, 4.2))
    region.plot(kind="bar", color="#4C78A8")
    plt.title("Sales by Region")
    plt.ylabel("Sales Amount")
    plt.xticks(rotation=0)
    plt.tight_layout()
    figs["region"] = FIG_DIR / "sales_by_region.png"
    plt.savefig(figs["region"], dpi=180)
    plt.close()

    category = df.groupby("Product Category")["Sales Amount"].sum().sort_values(ascending=False)
    plt.figure(figsize=(7, 4.2))
    category.plot(kind="bar", color="#72B7B2")
    plt.title("Sales by Product Category")
    plt.ylabel("Sales Amount")
    plt.xticks(rotation=35, ha="right")
    plt.tight_layout()
    figs["category"] = FIG_DIR / "sales_by_category.png"
    plt.savefig(figs["category"], dpi=180)
    plt.close()

    monthly = df.resample("ME", on="Order Date")["Sales Amount"].sum()
    plt.figure(figsize=(7, 4.2))
    monthly.plot(marker="o", color="#F58518")
    plt.title("Monthly Sales Trend")
    plt.ylabel("Sales Amount")
    plt.xlabel("Month")
    plt.tight_layout()
    figs["trend"] = FIG_DIR / "monthly_sales_trend.png"
    plt.savefig(figs["trend"], dpi=180)
    plt.close()

    return figs


def title_page(doc):
    for _ in range(2):
        para(doc)
    para(doc, "A REPORT OF SEMESTER INDUSTRIAL TRAINING", size=14, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)
    para(doc, "at", size=12, align=WD_ALIGN_PARAGRAPH.CENTER)
    para(doc, COMPANY, size=14, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)
    para(doc)
    para(doc, "SUBMITTED IN PARTIAL FULFILLMENT OF THE REQUIREMENT FOR THE AWARD", size=12, align=WD_ALIGN_PARAGRAPH.CENTER)
    para(doc, "OF THE DEGREE OF", size=12, align=WD_ALIGN_PARAGRAPH.CENTER)
    para(doc, "BACHELOR OF TECHNOLOGY", size=14, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)
    para(doc, f"({PROGRAM})", size=14, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)
    para(doc)
    para(doc, TRAINING_PERIOD.upper(), size=14, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)
    for _ in range(2):
        para(doc)
    para(doc, "PROJECT TITLE", size=12, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)
    para(doc, "ADVANCED SALES & CUSTOMER INTELLIGENCE PLATFORM", size=14, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)
    for _ in range(2):
        para(doc)
    para(doc, "SUBMITTED BY:", size=12, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)
    para(doc, f"NAME: {STUDENT_NAME}", size=12, align=WD_ALIGN_PARAGRAPH.CENTER)
    para(doc, f"UNIVERSITY ROLL NO.: {ROLL_NO}", size=12, align=WD_ALIGN_PARAGRAPH.CENTER)
    for _ in range(2):
        para(doc)
    para(doc, "DEPARTMENT OF COMPUTER SCIENCE AND ENGINEERING", size=12, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)
    para(doc, "I.K. GUJRAL PUNJAB TECHNICAL UNIVERSITY", size=14, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)


def front_matter(doc):
    doc.add_page_break()
    para(doc, "CERTIFICATE BY COMPANY / INDUSTRY / INSTITUTE", size=14, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)
    para(doc)
    para(doc, f"This is to certify that {STUDENT_NAME}, University Roll No. {ROLL_NO}, has successfully completed semester industrial training at {COMPANY} during {TRAINING_PERIOD}. The project titled \"Advanced Sales & Customer Intelligence Platform\" was carried out as part of the training work in the field of Data Science, Business Analytics and Web Application Development.")
    para(doc, "The candidate has worked on data cleaning, feature engineering, database loading, dashboard development, machine learning based analysis, forecasting, association rule mining and deployment support. The work presented in this report is an authentic record of the training work completed during the training period.")
    for _ in range(3):
        para(doc)
    para(doc, "Date: ____________                                Signature of Guide/Supervisor: ____________", spacing=1)
    para(doc, "Place: ___________                               Name and Seal: ___________________________", spacing=1)

    doc.add_page_break()
    para(doc, "I.K. GUJRAL PUNJAB TECHNICAL UNIVERSITY, KAPURTHALA", size=12, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)
    para(doc, "CANDIDATE'S DECLARATION", size=14, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)
    para(doc)
    para(doc, f"I, {STUDENT_NAME}, hereby declare that I have undertaken semester industrial training at {COMPANY} during the period {TRAINING_PERIOD} in partial fulfillment of the requirements for the award of degree of B.Tech. Computer Science and Engineering at I.K. Gujral Punjab Technical University, Kapurthala.")
    para(doc, "The work presented in this training report submitted to the Department of Computer Science and Engineering is an authentic record of the training work carried out by me. The report has not been submitted elsewhere for any other degree, diploma or certificate.")
    for _ in range(3):
        para(doc)
    para(doc, "Signature of the Student", spacing=1)
    para(doc)
    para(doc, "The semester industrial training Viva-Voce Examination of __________________ has been held on ____________ and accepted.", spacing=1)
    para(doc)
    para(doc, "Signature of Internal Examiner                         Signature of External Examiner", spacing=1)

    doc.add_page_break()
    para(doc, "ABSTRACT", size=14, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)
    para(doc)
    para(doc, "This report presents the design and development of an Advanced Sales & Customer Intelligence Platform, an enterprise-grade analytics dashboard created using Python, Streamlit and SQLite. The platform converts raw retail sales transactions into interactive business intelligence views, predictive forecasts, customer segments, product association rules, scenario simulations and SQL-based custom analytics.")
    para(doc, "The project begins with an Extract-Transform-Load pipeline that cleans a raw sales dataset, handles duplicates and missing values, creates temporal columns and enriches the data with derived business features such as profit margin, customer lifetime value, customer frequency, weekend flag, season, order value segment and customer tier. The enriched dataset is loaded into a local SQLite database for structured querying and dashboard consumption.")
    para(doc, "The dashboard provides six major workspaces: Executive Dashboard, Advanced Forecasting, Customer Analytics, Market Basket Analysis, What-If Simulator and SQL Analytics Workspace. It applies machine learning and statistical techniques including K-Means clustering, Random Forest regression, Prophet forecasting and Apriori association rule mining. The system demonstrates how modern data science tools can be integrated into a single decision-support platform for sales leaders, marketing teams, finance users and business executives.")
    para(doc, "Keywords: Business Analytics, Streamlit, SQLite, Machine Learning, Forecasting, Customer Segmentation, Market Basket Analysis, Python.")

    doc.add_page_break()
    para(doc, "ACKNOWLEDGEMENT", size=14, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)
    para(doc)
    para(doc, "I express my sincere gratitude to my training guide, faculty members, mentors and all individuals who supported me during the successful completion of this semester industrial training project. Their guidance helped me understand practical aspects of data analytics, software development, database management and machine learning implementation.")
    para(doc, "I am thankful to the Department of Computer Science and Engineering, I.K. Gujral Punjab Technical University, Kapurthala, for providing the academic framework and encouragement required to complete this work. I also acknowledge the open-source community for libraries such as pandas, NumPy, Streamlit, Plotly, scikit-learn, Prophet and mlxtend, which made the implementation of this analytics platform possible.")

    doc.add_page_break()
    para(doc, "ABOUT THE COMPANY / INDUSTRY / INSTITUTE", size=14, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)
    para(doc)
    para(doc, f"The semester industrial training was undertaken at {COMPANY}. During the training, the focus was on practical exposure to software tools, data science workflows and analytics application development. The work involved understanding business data, preparing datasets, building analytical models, creating interactive dashboards and documenting the complete solution.")
    para(doc, "The training environment helped connect classroom learning with real-world development practices. The project required programming, data cleaning, database handling, machine learning, visualization, documentation and deployment awareness.")


def contents(doc):
    doc.add_page_break()
    para(doc, "TABLE OF CONTENTS", size=14, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)
    rows = [
        ["Topic", "Page No."],
        ["Certificate by Company/Industry/Institute", "i"],
        ["Candidate's Declaration", "ii"],
        ["Abstract", "iii"],
        ["Acknowledgement", "iv"],
        ["About the Company / Industry / Institute", "v"],
        ["List of Figures", "vi"],
        ["List of Tables", "vii"],
        ["Definitions, Acronyms and Abbreviations", "viii"],
        ["CHAPTER 1 INTRODUCTION", "1"],
        ["CHAPTER 2 FIELD OF TRAINING", "10"],
        ["CHAPTER 3 TRAINING WORK UNDERTAKEN", "18"],
        ["CHAPTER 4 RESULTS AND DISCUSSIONS", "34"],
        ["CHAPTER 5 CONCLUSION AND FUTURE SCOPE", "48"],
        ["REFERENCES", "51"],
        ["APPENDIX", "53"],
    ]
    add_table(doc, rows, widths=[13.5, 3.0])

    doc.add_page_break()
    para(doc, "LIST OF FIGURES", size=14, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)
    rows = [
        ["Figure No.", "Title", "Page No."],
        ["Figure 1.1", "Overall data pipeline architecture", "4"],
        ["Figure 3.1", "Sales by region generated from enriched dataset", "27"],
        ["Figure 3.2", "Sales by product category generated from enriched dataset", "28"],
        ["Figure 4.1", "Monthly sales trend generated from enriched dataset", "39"],
    ]
    add_table(doc, rows, widths=[3.0, 10.5, 3.0])

    doc.add_page_break()
    para(doc, "LIST OF TABLES", size=14, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)
    rows = [
        ["Table No.", "Title", "Page No."],
        ["Table 1.1", "Project modules and purpose", "6"],
        ["Table 1.2", "Technology stack used in the project", "8"],
        ["Table 3.1", "Dataset profile after enrichment", "22"],
        ["Table 3.2", "Feature engineering details", "25"],
        ["Table 4.1", "Dashboard features and outputs", "36"],
        ["Table 4.2", "Key project outcomes", "45"],
    ]
    add_table(doc, rows, widths=[3.0, 10.5, 3.0])

    doc.add_page_break()
    para(doc, "DEFINITIONS, ACRONYMS AND ABBREVIATIONS", size=14, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)
    rows = [
        ["Term", "Meaning"],
        ["AOV", "Average Order Value"],
        ["CSV", "Comma-Separated Values"],
        ["EDA", "Exploratory Data Analysis"],
        ["ETL", "Extract, Transform and Load"],
        ["KPI", "Key Performance Indicator"],
        ["LTV", "Lifetime Value"],
        ["MAPE", "Mean Absolute Percentage Error"],
        ["ML", "Machine Learning"],
        ["RFM", "Recency, Frequency and Monetary analysis"],
        ["RMSE", "Root Mean Squared Error"],
        ["SQL", "Structured Query Language"],
    ]
    add_table(doc, rows, widths=[3.5, 13.0])


def chapter1(doc):
    chapter(doc, "CHAPTER 1 INTRODUCTION")
    heading(doc, "1.1 Background of the Project")
    para(doc, "In the modern business environment, organisations collect large volumes of transactional data through sales systems, customer interactions, payment channels and digital commerce platforms. However, the value of this data depends on how effectively it is cleaned, organised, analysed and presented to decision makers. Static spreadsheets and manually prepared monthly reports often fail to provide timely and interactive insights.")
    para(doc, "The Advanced Sales & Customer Intelligence Platform addresses this problem by converting raw sales records into an integrated analytics system. It combines data engineering, visual analytics, machine learning and database querying in a single Streamlit application. The system is designed around practical business questions such as which regions perform best, which product categories drive revenue, what future sales may look like, which customers are valuable and which products are likely to sell together.")
    para(doc, "The project is relevant because it demonstrates the complete lifecycle of a data analytics product: raw dataset ingestion, cleaning, feature engineering, database preparation, model building, dashboard design, user interaction and deployment readiness.")

    heading(doc, "1.2 Problem Statement")
    para(doc, "Retail and sales teams frequently face difficulty in transforming transaction-level data into meaningful business intelligence. Raw data may contain duplicates, missing values, inconsistent dates and columns that are unsuitable for direct analysis. Even after basic cleaning, business users need dashboards, forecasting models, customer segmentation, recommendation logic and flexible SQL access to answer detailed questions.")
    para(doc, "The problem addressed in this training work is the design and implementation of a Python-based analytics platform that can process sales data and provide interactive decision-support features for multiple business stakeholders.")

    heading(doc, "1.3 Objectives of the Training Work")
    for item in [
        "To design a complete data cleaning and preprocessing workflow for retail sales transactions.",
        "To engineer meaningful business features such as profit margin, customer lifetime value, customer tier and season.",
        "To load the enriched dataset into a SQLite database for structured query access.",
        "To build an interactive Streamlit dashboard with executive KPIs and visual analytics.",
        "To implement forecasting using Prophet and Random Forest models.",
        "To implement RFM based customer segmentation using K-Means clustering.",
        "To implement market basket analysis using Apriori association rule mining.",
        "To create a what-if simulator for price, volume and discount decisions.",
        "To provide a SQL analytics workspace for custom business analysis.",
    ]:
        bullet(doc, item)

    heading(doc, "1.4 Scope of the Project")
    para(doc, "The scope of the project includes sales data preparation, business intelligence dashboard development, machine learning based analysis and local deployment. The platform is suitable for academic demonstration, portfolio presentation and small-scale business analytics. It uses a local dataset and a local SQLite database, making it easy to run on a personal computer without cloud infrastructure.")
    para(doc, "The system does not currently include real-time API ingestion, role-based authentication, enterprise security policies or large-scale distributed storage. These features are considered future enhancements.")

    heading(doc, "1.5 Project Architecture")
    para(doc, "The architecture follows a simple but effective analytics pipeline. Raw sales data is cleaned, enriched, stored in SQLite and then consumed by the Streamlit dashboard. The dashboard applies filters, aggregations, visualisations and machine learning algorithms on top of the processed data.")
    para(doc, "Figure 1.1 Overall data pipeline architecture", size=10, italic=True, align=WD_ALIGN_PARAGRAPH.CENTER, spacing=1)
    para(doc, "Raw Sales CSV -> clean_sales_data.py -> cleaned_sales_dataset.csv -> enrich_sales_data.py -> enriched_sales_dataset.csv -> load_db.py -> sales.db -> app.py Streamlit Dashboard", align=WD_ALIGN_PARAGRAPH.CENTER)

    table_caption(doc, "Table 1.1 Project modules and purpose")
    add_table(doc, [
        ["Module", "Purpose"],
        ["clean_sales_data.py", "Removes duplicates, handles missing values, converts dates and creates Month and Year fields."],
        ["enrich_sales_data.py", "Adds business and analytical features such as profit margin, LTV, frequency, tiers and season."],
        ["load_db.py", "Loads the enriched dataset into SQLite tables named sales and sales_enriched."],
        ["app.py", "Provides the six-tab Streamlit analytics dashboard."],
        ["eda_sales.py", "Performs exploratory data analysis and creates static charts."],
        ["sales_prediction.py", "Provides a standalone linear regression forecasting experiment."],
    ], widths=[5.0, 11.5])

    table_caption(doc, "Table 1.2 Technology stack used in the project")
    add_table(doc, [
        ["Area", "Technology"],
        ["Programming Language", "Python"],
        ["User Interface", "Streamlit"],
        ["Data Processing", "pandas, NumPy"],
        ["Visualisation", "Plotly, Matplotlib, Seaborn"],
        ["Machine Learning", "scikit-learn"],
        ["Forecasting", "Prophet, Random Forest Regressor"],
        ["Association Mining", "mlxtend Apriori"],
        ["Database", "SQLite"],
        ["Report Export", "FPDF2"],
        ["Deployment", "Docker, Docker Compose, Windows Batch Script"],
    ], widths=[5.0, 11.5])


def chapter2(doc):
    chapter(doc, "CHAPTER 2 FIELD OF TRAINING")
    heading(doc, "2.1 Data Science and Business Analytics")
    para(doc, "The field of training for this project is Data Science and Business Analytics. Data Science involves collecting, cleaning, transforming, modelling and interpreting data to extract useful information. Business Analytics applies these techniques to business problems such as sales performance, customer retention, forecasting, pricing and product recommendations.")
    para(doc, "In this project, data science concepts are used practically through ETL processing, statistical aggregation, feature engineering, machine learning, visualisation and decision simulation.")

    heading(doc, "2.2 Data Engineering Concepts")
    para(doc, "Data engineering is the foundation of the platform. Before any dashboard or model can be trusted, the data must be prepared. The cleaning script removes duplicate Order IDs, drops rows with missing critical fields, fills missing customer and product names with Unknown, imputes missing profit values using the median and converts order dates into a datetime format.")
    para(doc, "The enrichment script then creates additional analytical columns that provide richer business context. These derived fields allow the dashboard to present more advanced KPIs than raw sales totals.")

    heading(doc, "2.3 Machine Learning Concepts Used")
    para(doc, "Machine learning is used in three major areas of the platform. First, K-Means clustering groups customers according to Recency, Frequency and Monetary value. Second, Random Forest regression forecasts future sales using time index and lag features. Third, Apriori association rule mining identifies products that are frequently purchased together by the same customer or in the same order.")
    para(doc, "These algorithms represent unsupervised learning, supervised learning and pattern mining. Their use in one project shows how multiple analytical methods can be combined to support different business questions.")

    heading(doc, "2.4 Time-Series Forecasting")
    para(doc, "Time-series forecasting is used to estimate future sales based on historical monthly revenue patterns. The dashboard compares Prophet, which models trends and seasonality, with Random Forest, which learns from lagged sales features. The app also calculates validation metrics such as MAPE and RMSE to help evaluate model performance.")

    heading(doc, "2.5 Dashboard and Visual Analytics")
    para(doc, "Dashboards convert processed data into interactive visuals. Streamlit provides an efficient way to build browser-based dashboards in Python. Plotly adds interactive charts such as bar charts, pie charts, line charts, 3D scatter plots and heatmaps. These visuals help non-technical users explore sales patterns without writing code.")


def chapter3(doc, stats, figs):
    chapter(doc, "CHAPTER 3 TRAINING WORK UNDERTAKEN")
    heading(doc, "3.1 Dataset Understanding")
    para(doc, f"The project uses a retail sales dataset that originally contains raw transaction records. After cleaning and enrichment, the final dataset contains {stats['rows']:,} records, {stats['orders']:,} unique orders, {stats['customers']:,} customers and {stats['products']:,} products. The dataset covers the years {stats['years']} across the regions {stats['regions']}.")
    table_caption(doc, "Table 3.1 Dataset profile after enrichment")
    add_table(doc, [
        ["Metric", "Value"],
        ["Total enriched records", f"{stats['rows']:,}"],
        ["Unique orders", f"{stats['orders']:,}"],
        ["Unique customers", f"{stats['customers']:,}"],
        ["Unique products", f"{stats['products']:,}"],
        ["Total sales", f"${stats['sales']:,.2f}"],
        ["Total profit", f"${stats['profit']:,.2f}"],
        ["Overall profit margin", f"{stats['margin']:.2f}%"],
        ["Average discount", f"{stats['avg_discount']:.2f}%"],
        ["Columns after enrichment", f"{stats['columns']}"],
    ], widths=[7.0, 9.5])

    heading(doc, "3.2 Data Cleaning Work")
    para(doc, "The first practical task was to prepare a reliable dataset. The cleaning script reads the raw CSV file, inspects the data shape and missing values, removes duplicate records based on Order ID, drops rows where critical fields are missing and fills less critical categorical values with Unknown.")
    for item in [
        "Order ID, Order Date, Region, Sales Amount and Quantity are treated as critical columns.",
        "Customer Name and Product Name are filled with Unknown when missing.",
        "Profit is filled with the median value to avoid bias from extreme values.",
        "Order Date is converted to datetime format for time-based analysis.",
        "Month and Year columns are created for grouping and filtering.",
    ]:
        bullet(doc, item)

    heading(doc, "3.3 Feature Engineering and Enrichment")
    para(doc, "Feature engineering expands the usefulness of the dataset. The enrichment script adds temporal, pricing, profitability, customer-level and product-level features. These features make it possible to calculate advanced KPIs, perform segmentation and build richer charts.")
    table_caption(doc, "Table 3.2 Feature engineering details")
    add_table(doc, [
        ["Feature Group", "New Features"],
        ["Temporal", "Quarter, Is_Weekend, Season, Day_of_Week"],
        ["Pricing", "Unit_Price, Discount_Amount, AOV_Segment"],
        ["Profitability", "Profit_Margin_Pct, Avg_Product_Margin_Pct"],
        ["Customer", "Customer_LTV, Customer_Frequency, Is_Repeat_Customer, Customer_Tier"],
        ["Product", "Product_Revenue_Rank"],
    ], widths=[4.0, 12.5])

    heading(doc, "3.4 Database Loading")
    para(doc, "The load_db.py script creates a local SQLite database named sales.db. It reads the enriched CSV file, converts column names with spaces into SQL-friendly underscore-separated names and writes the data into the sales table. It also creates a sales_enriched table containing the enriched fields. The database is used by the dashboard and the SQL workspace.")

    heading(doc, "3.5 Dashboard Development")
    para(doc, "The main application file app.py builds a six-tab Streamlit dashboard. Users can select either the local SQLite database or an uploaded CSV file as the data source. Global sidebar filters allow analysis by region, product category and year range. These filters are applied across the analytical views.")
    for item in [
        "Executive Dashboard shows core KPIs, visual charts, automated insights and download options.",
        "Advanced Forecasting compares Prophet and Random Forest models.",
        "Customer Analytics provides RFM clustering and cohort retention.",
        "Market Basket Analysis generates association rules and cross-sell suggestions.",
        "What-If Simulator models revenue, cost and profit under pricing, volume and discount changes.",
        "SQL Workspace allows custom SQLite queries and CSV export of results.",
    ]:
        bullet(doc, item)

    heading(doc, "3.6 Exploratory Visual Analysis")
    doc.add_picture(str(figs["region"]), width=Inches(5.7))
    caption(doc, "Figure 3.1 Sales by region generated from enriched dataset")
    doc.add_picture(str(figs["category"]), width=Inches(5.7))
    caption(doc, "Figure 3.2 Sales by product category generated from enriched dataset")

    heading(doc, "3.7 Deployment Support")
    para(doc, "The project includes setup.bat for Windows users and Docker files for containerized execution. The batch file creates a virtual environment, installs dependencies, runs cleaning and database loading, and starts the Streamlit dashboard. The Dockerfile uses Python 3.10 slim, installs dependencies and starts the app on port 8501 through Docker Compose.")


def chapter4(doc, figs):
    chapter(doc, "CHAPTER 4 RESULTS AND DISCUSSIONS")
    heading(doc, "4.1 Executive Dashboard Results")
    para(doc, "The Executive Dashboard provides an immediate business overview. It calculates total sales, total profit, total orders and average order value. When enriched data is available, it also displays average profit margin, average customer lifetime value, repeat customer percentage and weekend sales share.")
    table_caption(doc, "Table 4.1 Dashboard features and outputs")
    add_table(doc, [
        ["Dashboard Area", "Output"],
        ["KPI Cards", "Total sales, total profit, order count, average order value and enriched KPIs."],
        ["Regional Analysis", "Bar chart showing sales distribution by region."],
        ["Category Analysis", "Donut chart showing revenue share by product category."],
        ["Trend Analysis", "Monthly sales line chart with markers."],
        ["Product Analysis", "Top 10 products by sales shown in a data table."],
        ["Downloads", "Filtered CSV export and executive PDF report export."],
    ], widths=[4.5, 12.0])

    heading(doc, "4.2 Forecasting Results")
    para(doc, "The forecasting module aggregates filtered sales data monthly and validates models on the most recent period. Prophet is used for trend and seasonality based forecasting, while Random Forest uses Month_Index, Lag_1 and Lag_2 as features. The user can choose a forecast horizon between three and twelve months and can visualise Prophet, Random Forest or both together.")
    doc.add_picture(str(figs["trend"]), width=Inches(5.7))
    caption(doc, "Figure 4.1 Monthly sales trend generated from enriched dataset")

    heading(doc, "4.3 Customer Analytics Results")
    para(doc, "The Customer Analytics tab applies RFM analysis to group customers according to purchase recency, purchase frequency and total monetary value. The values are transformed using log scaling and then standardized before K-Means clustering. Users can select between three and six clusters depending on the number of customers available in the filtered dataset.")
    para(doc, "The module also includes cohort retention analysis. Customers are grouped by first purchase month and the heatmap shows the percentage of customers returning in later months. This helps identify whether newly acquired customers continue purchasing over time.")

    heading(doc, "4.4 Market Basket Analysis Results")
    para(doc, "Market basket analysis identifies relationships between products. The application supports customer-level analysis, which finds products bought by the same customer over time, and order-level analysis, which checks products purchased in the same order. Association rules are generated using Apriori and are sorted by lift. Higher lift indicates a stronger cross-sell relationship.")

    heading(doc, "4.5 What-If Simulation Results")
    para(doc, "The What-If Simulator allows business users to test the effect of price changes, volume changes and discount changes. The simulator separates total cost into 30 percent fixed cost and 70 percent variable cost. Revenue changes with price, volume and discount, while variable cost changes with volume and fixed cost remains unchanged. The result is a comparison of actual and simulated revenue, cost and profit.")

    heading(doc, "4.6 SQL Workspace Results")
    para(doc, "The SQL Analytics Workspace gives users direct access to the SQLite database. It includes a query editor, schema helper, templates and a CSV download option for query results. This feature is useful for advanced users who want custom analysis beyond the dashboard charts.")
    table_caption(doc, "Table 4.2 Key project outcomes")
    add_table(doc, [
        ["Outcome", "Description"],
        ["Integrated analytics platform", "Six interactive modules are available in one Streamlit application."],
        ["Automated ETL pipeline", "Raw data is cleaned, enriched and stored in SQLite."],
        ["Machine learning usage", "K-Means, Random Forest, Prophet and Apriori are applied to business problems."],
        ["Decision support", "Forecasting, segmentation, recommendations and simulation help business planning."],
        ["Deployment readiness", "Windows setup and Docker configuration are included."],
    ], widths=[4.5, 12.0])


def chapter5(doc):
    chapter(doc, "CHAPTER 5 CONCLUSION AND FUTURE SCOPE")
    heading(doc, "5.1 Conclusion")
    para(doc, "The semester industrial training project successfully demonstrates the development of a complete sales and customer intelligence platform using Python, Streamlit and SQLite. The system begins with raw transaction data and converts it into a structured, enriched and queryable dataset. It then presents business insights through interactive dashboards, machine learning models and decision-support tools.")
    para(doc, "The project helped strengthen practical knowledge of data cleaning, feature engineering, database loading, dashboard design, forecasting, clustering, association rule mining, scenario modelling and deployment preparation. It also shows how different tools can be combined into one cohesive analytics application that supports executives, sales teams, marketing teams, finance users and analysts.")

    heading(doc, "5.2 Future Scope")
    for item in [
        "Add real-time data ingestion from APIs, ERP systems or cloud databases.",
        "Add authentication and role-based access control for enterprise use.",
        "Introduce churn prediction and customer lifetime value forecasting.",
        "Use advanced forecasting models such as SARIMA, XGBoost or LSTM for comparison.",
        "Add automated alerts for abnormal sales drops, high discounts or negative profit patterns.",
        "Deploy the application on a cloud platform with scheduled data refresh.",
        "Improve report generation by adding automatic screenshots and scheduled PDF delivery.",
    ]:
        bullet(doc, item)


def references_appendix(doc):
    doc.add_page_break()
    para(doc, "REFERENCES", size=14, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)
    refs = [
        "W. McKinney, Python for Data Analysis, 2nd ed. Sebastopol, USA: O'Reilly Media, 2017.",
        "G. James, D. Witten, T. Hastie, and R. Tibshirani, An Introduction to Statistical Learning. New York, USA: Springer, 2013.",
        "S. J. Taylor and B. Letham, \"Forecasting at scale,\" The American Statistician, vol. 72, no. 1, pp. 37-45, 2018.",
        "R. Agrawal and R. Srikant, \"Fast algorithms for mining association rules,\" Proceedings of the 20th International Conference on Very Large Data Bases, Santiago, Chile, 1994, pp. 487-499.",
        "F. Pedregosa et al., \"Scikit-learn: Machine Learning in Python,\" Journal of Machine Learning Research, vol. 12, pp. 2825-2830, 2011.",
        "Streamlit Inc., \"Streamlit Documentation,\" Online. Available: https://docs.streamlit.io/.",
        "Plotly Technologies Inc., \"Plotly Python Documentation,\" Online. Available: https://plotly.com/python/.",
        "Python Software Foundation, \"Python Documentation,\" Online. Available: https://docs.python.org/3/.",
    ]
    for r in refs:
        number_item(doc, r)

    doc.add_page_break()
    para(doc, "APPENDIX", size=14, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)
    heading(doc, "A. Project File Structure", level=2)
    para(doc, "The major files in the project include app.py, clean_sales_data.py, enrich_sales_data.py, load_db.py, sales_prediction.py, eda_sales.py, requirements.txt, Dockerfile, docker-compose.yml, setup.bat, sales_sql_queries.sql, README.md, project-overview.md, business_insights.md and the dataset folder.")
    heading(doc, "B. Sample SQL Query", level=2)
    para(doc, "SELECT Region, ROUND(SUM(Sales_Amount), 2) AS Total_Sales, COUNT(*) AS Order_Count FROM sales GROUP BY Region ORDER BY Total_Sales DESC;", spacing=1)
    heading(doc, "C. Main Python Libraries", level=2)
    para(doc, "The implementation uses pandas, NumPy, scikit-learn, matplotlib, seaborn, Streamlit, Plotly, Prophet, mlxtend and fpdf2.")
    heading(doc, "D. Notes for Final Submission", level=2)
    for item in [
        "Replace placeholders for student name, roll number, company/institute and signatures before printing.",
        "Add actual deployment screenshots if required by the department.",
        "Export a PDF copy from Microsoft Word after final proofreading.",
        "Bind the final report according to the university hard-bound submission instructions.",
    ]:
        bullet(doc, item)


def main():
    df, stats = load_stats()
    figs = make_figures(df)
    doc = Document()
    configure_document(doc)
    title_page(doc)
    front_matter(doc)
    contents(doc)
    chapter1(doc)
    chapter2(doc)
    chapter3(doc, stats, figs)
    chapter4(doc, figs)
    chapter5(doc)
    references_appendix(doc)
    doc.core_properties.title = "Industrial Training Report - Sales Intelligence Platform"
    doc.core_properties.author = STUDENT_NAME
    doc.core_properties.subject = "Semester Industrial Training Report"
    doc.core_properties.keywords = "Streamlit, Python, Sales Analytics, Machine Learning, SQLite"
    doc.core_properties.created = datetime.now()
    doc.save(OUT)
    print(f"Generated: {OUT}")


if __name__ == "__main__":
    main()
